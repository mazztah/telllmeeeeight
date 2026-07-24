# dv.py – Universal Data Processor (FINAL + kompatibel mit main.py)
import logging
import os
import mimetypes
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
from pydub import AudioSegment
from docx import Document
import PyPDF2
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import json

logger = logging.getLogger(__name__)

def get_mime(file_path: str) -> str:
    mime, _ = mimetypes.guess_type(file_path)
    if mime:
        return mime
    ext = os.path.splitext(file_path)[1].lower()
    fallback = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.png': 'image/png',
        '.jpg': 'image/jpeg',
        '.mp3': 'audio/mpeg',
        '.wav': 'audio/wav',
        '.mp4': 'video/mp4',
        '.txt': 'text/plain',
        '.csv': 'text/csv',
        '.json': 'application/json',
    }
    return fallback.get(ext, "application/octet-stream")


# ====================== LESEN ======================
def _ocr_pdf_fallback(file_path: str, num_pages: int, max_pages: int = 10, dpi: int = 120,
                       min_chars_early_exit: int = 11000, progress: dict = None) -> str:
    """OCR-Fallback fuer PDFs ohne (verwertbaren) Text-Layer, z.B. eingescannte
    Bewerbungsmappen. Bewusst limitiert (max_pages, niedrige DPI), damit das
    auch auf schwacher/gedrosselter Free-Tier-CPU in vertretbarer Zeit
    durchlaeuft - die eigentlichen Lebenslauf-Seiten stehen bei Bewerbungs-
    mappen ohnehin fast immer am Anfang der Datei; nachfolgende Zeugnisse/
    Zertifikate sind fuer die reine Profil-Extraktion zweitrangig.

    NEU: Seite fuer Seite statt alles-auf-einmal, mit zwei Verbesserungen:
    - Early-Exit: sobald genug Text zusammengekommen ist (min_chars_early_exit),
      wird abgebrochen statt stur alle max_pages abzuarbeiten - spart auf
      langsamer Hardware oft die Haelfte der Zeit.
    - progress (optional, geteiltes dict): wird nach JEDER Seite aktualisiert.
      Bricht der Aufrufer wegen eines Gesamt-Timeouts ab, kann er trotzdem
      den bereits erkannten Text aus 'progress' weiterverwenden statt bei
      Null anzufangen ("in Haeppchen arbeiten" statt alles-oder-nichts).

    Bei fehlendem Tesseract/Sprachpaket wird sauber degradiert (kein Crash) -
    der Aufrufer behandelt eine leere Rueckgabe dann als "kein Text erkennbar".
    """
    if progress is None:
        progress = {}
    progress.setdefault("text", "")
    progress.setdefault("pages_done", 0)

    try:
        import pytesseract
    except Exception as exc:
        logger.warning("OCR-Fallback nicht verfuegbar (pytesseract fehlt): %s", exc)
        return ""

    # PyMuPDF (fitz) statt pdf2image/poppler fuer das Seiten-Rendering, da
    # pymupdf bereits eine Projekt-Abhaengigkeit ist (reines Python-Wheel,
    # kein zusaetzliches OS-Paket wie poppler-utils noetig) - portabler fuer
    # unterschiedliche Deployment-Umgebungen.
    try:
        import fitz
        from PIL import Image
        import io as _io
    except Exception as exc:
        logger.warning("OCR-Fallback nicht verfuegbar (pymupdf/PIL fehlt): %s", exc)
        return ""

    try:
        doc = fitz.open(file_path)
        pages_to_scan = min(num_pages, max_pages, len(doc))
    except Exception as exc:
        logger.warning("OCR-Fallback: PDF konnte nicht geoeffnet werden: %s", exc)
        return ""

    lang_attempts = ["deu+eng", "eng", None]
    texts = []
    try:
        for i in range(pages_to_scan):
            try:
                pix = doc[i].get_pixmap(dpi=dpi)
                img = Image.open(_io.BytesIO(pix.tobytes("png")))
            except Exception as exc:
                logger.warning("OCR-Fallback: Seite %d konnte nicht gerendert werden: %s", i + 1, exc)
                continue

            page_text = ""
            for lang in lang_attempts:
                try:
                    page_text = pytesseract.image_to_string(img, lang=lang) if lang else pytesseract.image_to_string(img)
                    break
                except Exception as exc:
                    logger.debug("OCR Sprachpaket '%s' fehlgeschlagen: %s", lang, exc)
                    continue

            texts.append(page_text)
            progress["pages_done"] = i + 1
            progress["text"] = "\n".join(texts).strip()

            if len(progress["text"]) >= min_chars_early_exit:
                logger.info("OCR-Fallback: Early-Exit nach %d/%d Seiten (%d Zeichen erreicht)",
                             i + 1, pages_to_scan, len(progress["text"]))
                break
    finally:
        doc.close()

    combined = progress["text"]
    if combined:
        logger.info("OCR-Fallback erfolgreich: %d Seiten gescannt, %d Zeichen erkannt",
                     progress["pages_done"], len(combined))
    return combined


def extract_content(file_path: str, max_chars: int = 12000, ocr_progress: dict = None) -> str:
    mime = get_mime(file_path)
    name = os.path.basename(file_path)

    try:
        if mime.startswith("image/"):
            img = Image.open(file_path)
            return f"📸 Bild '{name}': {img.format} | Größe: {img.size} | Modus: {img.mode}"

        elif mime.startswith("audio/"):
            audio = AudioSegment.from_file(file_path)
            return f"🎵 Audio '{name}': {len(audio)/1000:.1f}s | Kanäle: {audio.channels} | {audio.frame_rate}Hz"

        elif mime.startswith("video/"):
            return f"🎥 Video '{name}': Datei erkannt"

        elif mime == "application/pdf":
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = "".join(p.extract_text() or "" for p in reader.pages)
            num_pages = len(reader.pages)

            # NEU: Viele "Bewerbungsmappen" sind eingescannte PDFs OHNE
            # Text-Layer (nur Bilder). PyPDF2 liefert dann exakt 0 (oder fast
            # 0) Zeichen zurueck. Ohne diesen Fallback wurde bisher trotzdem
            # ein (praktisch leerer) "Extrakt" an das LLM geschickt, das
            # daraufhin ein komplett erfundenes Profil halluziniert hat -
            # OHNE dass das fuer den Nutzer erkennbar war. Fallback: OCR ueber
            # die ersten Seiten (mehr Seiten waeren auf schwacher Hardware zu
            # langsam; die eigentlichen Lebenslauf-Seiten stehen ohnehin am
            # Anfang der Datei, Zeugnisse/Zertifikate danach sind sekundaer).
            if len(text.strip()) < 200:
                ocr_text = _ocr_pdf_fallback(file_path, num_pages, progress=ocr_progress)
                if ocr_text:
                    text = ocr_text
                    return (f"📄 PDF '{name}' ({num_pages} Seiten, per OCR erkannt "
                            f"- Scan ohne Text-Layer):\n{text[:max_chars]} [...]")

            return f"📄 PDF '{name}' ({num_pages} Seiten):\n{text[:max_chars]} [...]"

        elif "wordprocessingml" in mime or mime.endswith("document"):
            doc = Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return f"📝 Word '{name}':\n{text[:max_chars]} [...]"

        elif mime in ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "text/csv"):
            df = pd.read_excel(file_path) if "excel" in mime else pd.read_csv(file_path)
            return f"📊 {mime.split('/')[-1].upper()} '{name}' ({df.shape[0]} Zeilen):\n{df.head(10).to_string()}"

        elif mime == "application/json":
            with open(file_path, encoding="utf-8") as f:
                data = json.load(f)
            return f"📋 JSON '{name}':\n{json.dumps(data, indent=2, ensure_ascii=False)[:max_chars]}"

        elif mime.startswith("text/"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read(max_chars)
            return f"📜 Text '{name}':\n{text}"

        return f"❓ Unbekanntes Format '{name}' (MIME: {mime})"

    except Exception as e:
        logger.error(f"extract_content Fehler bei {name}: {e}")
        return f"💥 Extract-Fehler bei {name}: {str(e)[:200]}"


# ====================== ALTE FUNKTIONEN (für main.py Kompatibilität) ======================
def resize_image(file_path: str, size: tuple = (800, 800)) -> BytesIO:
    img = Image.open(file_path).resize(size, Image.Resampling.LANCZOS)
    buffer = BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def trim_audio(file_path: str, start: float = 0.0, end: float | None = None) -> BytesIO:
    audio = AudioSegment.from_file(file_path)
    if end is None:
        end = len(audio) / 1000
    trimmed = audio[int(start * 1000):int(end * 1000)]
    buffer = BytesIO()
    trimmed.export(buffer, format="mp3")
    buffer.seek(0)
    return buffer


def save_locally(buffer: BytesIO, filename: str, folder: str = "/tmp/processed") -> str:
    os.makedirs(folder, exist_ok=True)
    path = os.path.join(folder, filename)
    with open(path, "wb") as f:
        f.write(buffer.getvalue())
    return path


# ====================== NEUE ERSTELL-FUNKTIONEN ======================
def create_pdf_from_text(text: str, title: str = "Generated.pdf") -> BytesIO:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    y = 750
    for line in text.split("\n"):
        if y < 50:
            c.showPage()
            y = 750
        c.drawString(50, y, line[:90])
        y -= 15
    c.save()
    buffer.seek(0)
    return buffer


def create_docx_from_text(text: str, title: str = "Generated.docx") -> BytesIO:
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_excel_from_data(data: list[list], columns: list[str], title: str = "Generated.xlsx") -> BytesIO:
    """Erstellt ein formatiertes XLSX mit Hyperlinks für URL-Spalten."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        ws = wb.active
        ws.title = "Ergebnisse"
        ws.freeze_panes = "A2"  # Header einfrieren

        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill("solid", fgColor="6D28D9")
        header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        url_font = Font(color="0563C1", underline="single", size=10)
        cell_align = Alignment(vertical="top", wrap_text=True)
        thin = Side(style="thin", color="E5E7EB")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        # Spaltenbreiten je nach Typ
        col_widths = {
            "URL": 55, "Beschreibung (Snippet)": 55, "Titel": 38,
            "Firma": 28, "Ort": 22, "Job-ID": 14,
        }

        # Header-Zeile
        for ci, col_name in enumerate(columns, 1):
            cell = ws.cell(row=1, column=ci, value=col_name)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = border
            letter = get_column_letter(ci)
            ws.column_dimensions[letter].width = col_widths.get(col_name, 20)

        ws.row_dimensions[1].height = 22

        # Datenzeilen
        for ri, row in enumerate(data, 2):
            for ci, val in enumerate(row, 1):
                col_name = columns[ci - 1] if ci - 1 < len(columns) else ""
                str_val = str(val) if val is not None else ""
                cell = ws.cell(row=ri, column=ci, value=str_val)
                cell.alignment = cell_align
                cell.border = border
                # URL → Hyperlink
                if col_name == "URL" and str_val.startswith("http"):
                    cell.hyperlink = str_val
                    cell.value = str_val
                    cell.font = url_font
                else:
                    cell.font = Font(size=10)
                # Zebra-Streifen
                if ri % 2 == 0:
                    cell.fill = PatternFill("solid", fgColor="F5F3FF")

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer

    except ImportError:
        # Fallback auf pandas falls openpyxl fehlt
        df = pd.DataFrame(data, columns=columns)
        buffer = BytesIO()
        df.to_excel(buffer, index=False, engine="openpyxl")
        buffer.seek(0)
        return buffer



def create_jobqueen_excel(jobs: list, queries: list = None, export_date: str = "") -> BytesIO:
    """Erstellt professionelles JobQueen-Export-XLSX ohne Emoji-Sonderzeichen."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        import datetime as _dt
        from collections import Counter
    except ImportError as _e:
        logger.error("openpyxl nicht installiert: %s", _e)
        raise

    wb   = Workbook()
    thin = Side(style="thin",   color="C4B5FD")
    med  = Side(style="medium", color="7C3AED")
    brd  = Border(left=thin, right=thin, top=thin, bottom=thin)
    brdH = Border(left=med,  right=med,  top=med,  bottom=med)

    # aRGB Farben (FF = volle Deckkraft, kein # Praefix)
    COL_PURPLE = "FF6D28D9"
    COL_PURLT  = "FFEDE9FE"
    COL_GOLD   = "FFD4AF37"
    COL_WHITE  = "FFFFFFFF"
    COL_DARK   = "FF1E1B4B"
    COL_DARK2  = "FF312E81"
    COL_GRAY   = "FF6B7280"
    COL_GREEN  = "FF065F46"
    COL_GBGG   = "FFD1FAE5"
    COL_LINK   = "FF0563C1"

    def _H(ws, r, c, v, bg=None, fg=None, sz=11, bold=True, ctr=True):
        bg = bg or COL_PURPLE
        fg = fg or COL_WHITE
        x = ws.cell(row=r, column=c, value=v)
        x.font      = Font(bold=bold, color=fg, size=sz, name="Calibri")
        x.fill      = PatternFill("solid", fgColor=bg)
        x.alignment = Alignment(
            horizontal="center" if ctr else "left",
            vertical="center", wrap_text=True)
        x.border = brdH
        return x

    def _D(ws, r, c, v, bold=False, color=None, bg=None, url=None, ctr=False, sz=10):
        color = color or COL_DARK
        s = str(v) if v is not None else ""
        x = ws.cell(row=r, column=c, value=s)
        x.font      = Font(bold=bold, color=color, size=sz, name="Calibri",
                           underline="single" if url else None)
        x.alignment = Alignment(
            horizontal="center" if ctr else "left",
            vertical="top", wrap_text=True)
        x.border = brd
        if bg:
            x.fill = PatternFill("solid", fgColor=bg)
        if url and str(url).startswith("http"):
            x.hyperlink = str(url)
            x.value     = str(v) if v else str(url)
        return x

    # ── Sheet 1: Zusammenfassung ──────────────────────────────────────────
    ws1 = wb.active
    ws1.title = "Zusammenfassung"
    ws1.sheet_view.showGridLines = False

    ws1.merge_cells("A1:F1")
    c = ws1["A1"]
    c.value     = "JobQueen - Stellensuche Export"
    c.font      = Font(bold=True, color=COL_GOLD, size=18, name="Calibri")
    c.fill      = PatternFill("solid", fgColor=COL_DARK)
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws1.row_dimensions[1].height = 38

    ws1.merge_cells("A2:F2")
    c = ws1["A2"]
    ed = export_date or _dt.datetime.now().strftime("%d.%m.%Y %H:%M")
    c.value     = "Exportiert am {}   -   {} Stelle(n) ausgewaehlt".format(ed, len(jobs))
    c.font      = Font(color=COL_WHITE, size=10, italic=True, name="Calibri")
    c.fill      = PatternFill("solid", fgColor=COL_DARK2)
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws1.row_dimensions[2].height = 20
    ws1.row_dimensions[3].height = 8

    _H(ws1, 4, 1, "Suchanfrage", sz=10)
    _H(ws1, 4, 2, "Plattform",   sz=10)
    _H(ws1, 4, 3, "Stellen",     sz=10)

    src_counts = Counter(j.get("source") or "-" for j in jobs)
    qlist      = list(dict.fromkeys([q for q in (queries or []) if q]))
    n_rows     = max(len(qlist), len(src_counts), 1)

    for i in range(n_rows):
        bg = COL_PURLT if i % 2 == 0 else COL_WHITE
        _D(ws1, 5 + i, 1, qlist[i] if i < len(qlist) else "", bg=bg)
    for i, (src, cnt) in enumerate(src_counts.items()):
        bg = COL_PURLT if i % 2 == 0 else COL_WHITE
        _D(ws1, 5 + i, 2, src,      bg=bg)
        _D(ws1, 5 + i, 3, str(cnt), bg=bg, ctr=True)

    tr = 5 + n_rows + 1
    ws1.merge_cells("A{}:B{}".format(tr, tr))
    _H(ws1, tr, 1, "Gesamt: {} Stellen ausgewaehlt".format(len(jobs)),
       bg=COL_GREEN, sz=11)
    _D(ws1, tr, 3, "", bg=COL_GBGG)

    for col, w in zip("ABCDEF", [44, 28, 14, 10, 10, 10]):
        ws1.column_dimensions[col].width = w

    # ── Sheet 2: Ausgewaehlte Stellen ─────────────────────────────────────
    ws2 = wb.create_sheet("Ausgewaehlte Stellen")
    ws2.sheet_view.showGridLines = False
    ws2.freeze_panes = "A3"

    COLS = [
        ("Nr.",            6),
        ("Jobtitel",       38),
        ("Firma",          28),
        ("Ort",            20),
        ("Quelle",         20),
        ("Datum",          16),
        ("Link zur Stelle", 52),
        ("Kurzbeschreibung", 54),
    ]
    ncols = len(COLS)

    ws2.merge_cells("A1:{}1".format(get_column_letter(ncols)))
    c = ws2["A1"]
    c.value     = "JobQueen - {} ausgewaehlte Stellen".format(len(jobs))
    c.font      = Font(bold=True, color=COL_GOLD, size=14, name="Calibri")
    c.fill      = PatternFill("solid", fgColor=COL_DARK)
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws2.row_dimensions[1].height = 30

    for ci, (colname, w) in enumerate(COLS, 1):
        _H(ws2, 2, ci, colname, sz=10)
        ws2.column_dimensions[get_column_letter(ci)].width = w
    ws2.row_dimensions[2].height = 22

    for ri, j in enumerate(jobs, 3):
        bg  = COL_PURLT if (ri - 3) % 2 == 0 else COL_WHITE
        url = j.get("url") or ""
        raw_d = j.get("date") or ""
        try:
            date_str = _dt.datetime.fromisoformat(
                raw_d.split("T")[0]).strftime("%d.%m.%Y")
        except Exception:
            date_str = raw_d

        _D(ws2, ri, 1, ri - 2,
           bg=bg, bold=True, color=COL_PURPLE, ctr=True, sz=10)
        _D(ws2, ri, 2, j.get("title") or "",   bg=bg, bold=True, sz=10)
        _D(ws2, ri, 3, j.get("company") or "", bg=bg, sz=10)
        _D(ws2, ri, 4, j.get("location") or "", bg=bg, sz=10)
        _D(ws2, ri, 5, j.get("source") or "",  bg=bg, color=COL_GRAY, sz=9)
        _D(ws2, ri, 6, date_str,                bg=bg, ctr=True, sz=10)
        _D(ws2, ri, 7, "Zur Stelle oeffnen",
           bg=bg, url=url, color=COL_LINK, bold=True, sz=10)
        _D(ws2, ri, 8,
           (j.get("description_snippet") or "")[:300],
           bg=bg, color=COL_GRAY, sz=9)
        ws2.row_dimensions[ri].height = 44

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf



def create_jobqueen_pdf(jobs: list, queries: list = None, export_date: str = "") -> "BytesIO":
    """Erstellt professionelles JobQueen-PDF mit reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                        HRFlowable, KeepTogether)
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        buf = BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                rightMargin=18*mm, leftMargin=18*mm,
                                topMargin=18*mm, bottomMargin=18*mm)

        C_PURPLE = colors.HexColor("#6D28D9")
        C_DARK   = colors.HexColor("#1E1B4B")
        C_GRAY   = colors.HexColor("#6B7280")
        C_LIGHT  = colors.HexColor("#EDE9FE")
        C_GOLD   = colors.HexColor("#D4AF37")

        st_title = ParagraphStyle("jq_title", fontSize=22, fontName="Helvetica-Bold",
                                   textColor=C_PURPLE, alignment=TA_CENTER, spaceAfter=4)
        st_sub   = ParagraphStyle("jq_sub",   fontSize=10, fontName="Helvetica",
                                   textColor=C_GRAY,   alignment=TA_CENTER, spaceAfter=14)
        st_num   = ParagraphStyle("jq_num",   fontSize=9,  fontName="Helvetica-Bold",
                                   textColor=C_PURPLE)
        st_jt    = ParagraphStyle("jq_jt",    fontSize=13, fontName="Helvetica-Bold",
                                   textColor=C_DARK,   spaceAfter=3, spaceBefore=6)
        st_meta  = ParagraphStyle("jq_meta",  fontSize=9,  fontName="Helvetica",
                                   textColor=C_GRAY,   spaceAfter=2)
        st_url   = ParagraphStyle("jq_url",   fontSize=8,  fontName="Helvetica",
                                   textColor=colors.blue, spaceAfter=3)
        st_desc  = ParagraphStyle("jq_desc",  fontSize=9,  fontName="Helvetica",
                                   textColor=C_GRAY,   spaceAfter=6, leading=12)

        import datetime as _dt
        ed = export_date or _dt.datetime.now().strftime("%d.%m.%Y %H:%M")
        qstr = ", ".join(q for q in (queries or []) if q) or "-"

        story = [
            Paragraph("JobQueen - Stellensuche Export", st_title),
            Paragraph(f"Exportiert am {ed}  |  {len(jobs)} Stelle(n)  |  Suche: {qstr[:80]}", st_sub),
            HRFlowable(width="100%", thickness=2, color=C_PURPLE, spaceAfter=10),
        ]

        for i, j in enumerate(jobs, 1):
            raw_d = j.get("date") or ""
            try:
                import datetime as _dt2
                date_str = _dt2.datetime.fromisoformat(raw_d.split("T")[0]).strftime("%d.%m.%Y")
            except Exception:
                date_str = raw_d

            url     = j.get("url") or ""
            title   = (j.get("title") or "Unbekannte Stelle").replace("<", "&lt;").replace(">", "&gt;")
            company = (j.get("company") or "-").replace("<", "&lt;")
            loc     = (j.get("location") or "-").replace("<", "&lt;")
            src     = (j.get("source") or "").replace("<", "&lt;")
            desc    = (j.get("description_snippet") or "")[:280].replace("<", "&lt;")

            block = []
            block.append(Paragraph(f"[{i}]  {title}", st_jt))
            meta_parts = [f"Firma: {company}", f"Ort: {loc}"]
            if src:   meta_parts.append(f"Quelle: {src}")
            if date_str: meta_parts.append(f"Datum: {date_str}")
            block.append(Paragraph("  |  ".join(meta_parts), st_meta))
            if url:
                safe_url = url[:100]
                block.append(Paragraph(f"Link: {safe_url}", st_url))
            if desc:
                block.append(Paragraph(desc, st_desc))
            block.append(HRFlowable(width="100%", thickness=0.5,
                                    color=C_LIGHT, spaceAfter=4))
            story.append(KeepTogether(block))

        doc.build(story)
        buf.seek(0)
        return buf
    except Exception as _e:
        logger.error("create_jobqueen_pdf Fehler: %s", _e)
        # Plain-text fallback
        lines = ["JobQueen – Stellensuche Export", "=" * 40, ""]
        for i, j in enumerate(jobs, 1):
            lines.append(f"{i}. {j.get('title','?')} – {j.get('company','?')}")
            if j.get("url"): lines.append(f"   {j.get('url')}")
            lines.append("")
        text = "\n".join(lines)
        from reportlab.pdfgen import canvas as _c
        from reportlab.lib.pagesizes import letter as _letter
        buf = BytesIO()
        c = _c.Canvas(buf, pagesize=_letter)
        y = 750
        for line in text.split("\n"):
            if y < 50: c.showPage(); y = 750
            c.drawString(50, y, line[:90]); y -= 14
        c.save(); buf.seek(0)
        return buf


def create_jobqueen_markdown(jobs: list, queries: list = None, export_date: str = "") -> str:
    """Erstellt Markdown-Bericht der Stellensuche."""
    import datetime as _dt
    ed  = export_date or _dt.datetime.now().strftime("%d.%m.%Y %H:%M")
    qstr = ", ".join(q for q in (queries or []) if q) or "-"
    lines = [
        "# JobQueen – Stellensuche Export",
        "",
        f"| Feld | Wert |",
        f"|------|------|",
        f"| Exportiert am | {ed} |",
        f"| Stellen gesamt | {len(jobs)} |",
        f"| Suchanfragen | {qstr[:120]} |",
        "",
        "---",
        "",
    ]
    for i, j in enumerate(jobs, 1):
        raw_d = j.get("date") or ""
        try:
            import datetime as _dt2
            date_str = _dt2.datetime.fromisoformat(raw_d.split("T")[0]).strftime("%d.%m.%Y")
        except Exception:
            date_str = raw_d

        title   = j.get("title")   or "Unbekannte Stelle"
        company = j.get("company") or "-"
        loc     = j.get("location") or "-"
        src     = j.get("source")  or ""
        url     = j.get("url")     or ""
        desc    = (j.get("description_snippet") or "")[:300]

        lines += [
            f"## {i}. {title}",
            "",
            f"| | |",
            f"|---|---|",
            f"| **Firma** | {company} |",
            f"| **Ort** | {loc} |",
        ]
        if src:      lines.append(f"| **Quelle** | {src} |")
        if date_str: lines.append(f"| **Datum** | {date_str} |")
        if url:      lines.append(f"| **Link** | [{url[:60]}]({url}) |")
        lines.append("")
        if desc:
            lines += [f"> {desc}", ""]
        lines += ["---", ""]
    return "\n".join(lines)


def create_chart_from_df(df: pd.DataFrame, title: str = "Chart") -> BytesIO:
    img = Image.new("RGB", (900, 600), "#2a0044")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 24)
    except:
        font = ImageFont.load_default()

    draw.text((30, 30), title, fill="#ffccff", font=font)

    if len(df.columns) > 1:
        values = df.iloc[:, 1].head(8)
        max_val = max(values) if len(values) > 0 else 1
        for i, val in enumerate(values):
            x = 80 + i * 90
            height = int(380 * (val / max_val))
            draw.rectangle([x, 500 - height, x + 65, 500], fill="#ff66cc")
            draw.text((x + 10, 510), str(val), fill="white", font=font)

    buffer = BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def create_cv_excel(profile: dict) -> BytesIO:
    """Exportiert CV-Analyse-Ergebnisse als formatiertes XLSX."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    thin = Side(style="thin", color="D1D5DB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def hdr_cell(ws, row, col, value, fg="6D28D9"):
        c = ws.cell(row=row, column=col, value=value)
        c.font = Font(bold=True, color="FFFFFF", size=10)
        c.fill = PatternFill("solid", fgColor=fg)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.border = border
        return c

    def data_cell(ws, row, col, value, bold=False, color=None):
        c = ws.cell(row=row, column=col, value=str(value) if value is not None else "")
        c.font = Font(bold=bold, color=color or "1F2937", size=10)
        c.alignment = Alignment(vertical="top", wrap_text=True)
        c.border = border
        if row % 2 == 0:
            c.fill = PatternFill("solid", fgColor="F5F3FF")
        return c

    # ── Sheet 1: Übersicht ──────────────────────────────────────────
    ws1 = wb.active
    ws1.title = "Übersicht"
    ws1.freeze_panes = "A2"
    ws1.column_dimensions["A"].width = 28
    ws1.column_dimensions["B"].width = 55

    hdr_cell(ws1, 1, 1, "Feld")
    hdr_cell(ws1, 1, 2, "Wert")

    rows_overview = [
        ("Name",            profile.get("name") or "—"),
        ("Berufserfahrung", f"{profile.get('experience_years', 0)} Jahre {profile.get('experience_months', 0)} Monate"),
        ("Skills",         " · ".join(profile.get("skills") or []) or "—"),
        ("Sprachen",       " · ".join(profile.get("languages") or []) or "—"),
    ]
    for ri, (k, v) in enumerate(rows_overview, 2):
        data_cell(ws1, ri, 1, k, bold=True)
        data_cell(ws1, ri, 2, v)

    # ── Sheet 2: Stärken ───────────────────────────────────────────
    ws2 = wb.create_sheet("Stärken")
    ws2.freeze_panes = "A2"
    ws2.column_dimensions["A"].width = 30
    ws2.column_dimensions["B"].width = 50
    ws2.column_dimensions["C"].width = 45

    hdr_cell(ws2, 1, 1, "Stärke")
    hdr_cell(ws2, 1, 2, "Beleg aus dem Lebenslauf")
    hdr_cell(ws2, 1, 3, "Relevanz")

    for ri, s in enumerate(profile.get("strengths") or [], 2):
        data_cell(ws2, ri, 1, s.get("strength") or "", bold=True)
        data_cell(ws2, ri, 2, s.get("evidence") or "")
        data_cell(ws2, ri, 3, s.get("relevance") or "")

    # ── Sheet 3: Empfohlene Jobtitel ───────────────────────────────
    ws3 = wb.create_sheet("Empfohlene Jobtitel")
    ws3.freeze_panes = "A2"
    ws3.column_dimensions["A"].width = 38
    ws3.column_dimensions["B"].width = 60

    hdr_cell(ws3, 1, 1, "Jobtitel")
    hdr_cell(ws3, 1, 2, "Begründung")

    for ri, jt in enumerate(profile.get("suggested_job_titles") or [], 2):
        data_cell(ws3, ri, 1, jt.get("title") or "", bold=True)
        data_cell(ws3, ri, 2, jt.get("reason") or "")

    # ── Sheet 4: Berufserfahrung ────────────────────────────────────
    ws4 = wb.create_sheet("Berufserfahrung")
    ws4.freeze_panes = "A2"
    ws4.column_dimensions["A"].width = 32
    ws4.column_dimensions["B"].width = 30
    ws4.column_dimensions["C"].width = 14
    ws4.column_dimensions["D"].width = 14
    ws4.column_dimensions["E"].width = 12

    hdr_cell(ws4, 1, 1, "Position")
    hdr_cell(ws4, 1, 2, "Unternehmen")
    hdr_cell(ws4, 1, 3, "Von")
    hdr_cell(ws4, 1, 4, "Bis")
    hdr_cell(ws4, 1, 5, "Monate")

    roles = (profile.get("experience_details") or {}).get("roles") or []
    for ri, role in enumerate(roles, 2):
        data_cell(ws4, ri, 1, role.get("title") or "", bold=True)
        data_cell(ws4, ri, 2, role.get("company") or "")
        data_cell(ws4, ri, 3, role.get("start") or "")
        data_cell(ws4, ri, 4, role.get("end") or "")
        data_cell(ws4, ri, 5, role.get("months") or "")

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════════════════════════
# BEWERBUNGSANSCHREIBEN – professionell formatiertes DOCX + PDF
# ═══════════════════════════════════════════════════════════════════════════
#
# `letter` erwartet folgendes Dict-Schema (siehe main.py /coverletter/draft):
#   {
#     "subject": "Bewerbung als ...",
#     "salutation": "Sehr geehrte Damen und Herren,",
#     "paragraphs": ["Absatz 1 ...", "Absatz 2 ...", ...],
#     "closing": "Mit freundlichen Gruessen",
#     "signature_name": "Max Mustermann",
#   }
#
# `sender` (optional, alles darf leer/None sein):
#   {"name": "...", "address": "...", "email": "...", "phone": "...",
#    "website": "https://...", "linkedin": "https://...", "xing": "https://..."}
#
# `recipient` (optional):
#   {"company": "...", "attention": "...", "address": "..."}


def _cl_field(d: dict, key: str, default: str = "") -> str:
    val = (d or {}).get(key)
    return str(val).strip() if val else default


def _cl_norm_url(url: str) -> str:
    """Stellt sicher, dass ein Link ein gueltiges Schema hat (fuer klickbare
    Hyperlinks in DOCX/PDF), ohne den Nutzer mit Details zu belasten."""
    url = (url or "").strip()
    if not url:
        return ""
    if not url.startswith(("http://", "https://", "mailto:", "tel:")):
        url = "https://" + url
    return url


def _docx_add_hyperlink(paragraph, url: str, text: str, size_pt: float = 9.0):
    """python-docx bietet keine eingebaute Hyperlink-API - dieser Helfer fuegt
    einen echten klickbaren Link per OOXML-Relationship ein (Standard-Technik,
    da python-docx selbst keine add_hyperlink()-Methode hat)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rpr.append(sz)

    new_run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def create_cover_letter_docx(letter: dict, sender: dict = None, recipient: dict = None,
                              date_str: str = "") -> BytesIO:
    """Erstellt ein professionell formatiertes Bewerbungsanschreiben als DOCX
    (DIN-5008-angelehntes Layout: Absenderzeile, Datum rechtsbuendig, grosse
    Betreffzeile, Anrede, Fliesstext-Absaetze, Grussformel, klickbare
    Kontakt-Fusszeile mit Telefon/E-Mail/Website/LinkedIn/Xing)."""
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    sender = sender or {}
    recipient = recipient or {}

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    s_name = _cl_field(sender, "name")
    s_addr = _cl_field(sender, "address")
    s_mail = _cl_field(sender, "email")
    s_phone = _cl_field(sender, "phone")
    if s_name or s_addr or s_mail or s_phone:
        for line in filter(None, [s_name, s_addr, s_mail, s_phone]):
            p = doc.add_paragraph(line)
            p.runs[0].font.size = Pt(9.5)
            p.paragraph_format.space_after = Pt(0)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    r_company = _cl_field(recipient, "company")
    r_att = _cl_field(recipient, "attention")
    r_addr = _cl_field(recipient, "address")

    date_p = doc.add_paragraph(date_str)
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.paragraph_format.space_after = Pt(16)

    if r_company or r_att or r_addr:
        for line in filter(None, [r_company, r_att, r_addr]):
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(0)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # Betreffzeile: deutlich groesser + fett (vom Nutzer gewuenscht), etwas
    # mehr Luft davor/danach fuer bessere visuelle Trennung.
    subject = _cl_field(letter, "subject", "Bewerbung")
    subj_p = doc.add_paragraph()
    subj_p.paragraph_format.space_before = Pt(6)
    subj_p.paragraph_format.space_after = Pt(18)
    subj_run = subj_p.add_run(subject)
    subj_run.bold = True
    subj_run.font.size = Pt(14)

    salutation = _cl_field(letter, "salutation", "Sehr geehrte Damen und Herren,")
    sal_p = doc.add_paragraph(salutation)
    sal_p.paragraph_format.space_after = Pt(12)

    paragraphs = letter.get("paragraphs") or []
    if isinstance(paragraphs, str):
        paragraphs = [p.strip() for p in paragraphs.split("\n\n") if p.strip()]
    for para_text in paragraphs:
        p = doc.add_paragraph(para_text)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.25

    closing = _cl_field(letter, "closing", "Mit freundlichen Gruessen")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    close_p = doc.add_paragraph(closing)
    close_p.paragraph_format.space_after = Pt(36)

    sig_name = _cl_field(letter, "signature_name", s_name)
    if sig_name:
        doc.add_paragraph(sig_name)

    # ── Kontakt-Fusszeile: Telefon/E-Mail/Website/LinkedIn/Xing, klickbar ──
    s_web = _cl_norm_url(_cl_field(sender, "website"))
    s_linkedin = _cl_norm_url(_cl_field(sender, "linkedin"))
    s_xing = _cl_norm_url(_cl_field(sender, "xing"))
    if s_phone or s_mail or s_web or s_linkedin or s_xing:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(4)

        # duenne Trennlinie oberhalb der Fusszeile
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _Oxml
        pPr = fp._p.get_or_add_pPr()
        pBdr = _Oxml("w:pBdr")
        top = _Oxml("w:top")
        top.set(_qn("w:val"), "single")
        top.set(_qn("w:sz"), "4")
        top.set(_qn("w:space"), "6")
        top.set(_qn("w:color"), "CCCCCC")
        pBdr.append(top)
        pPr.append(pBdr)

        segments = []
        if s_phone:
            segments.append(("text", s_phone))
        if s_mail:
            segments.append(("link", f"mailto:{s_mail}", s_mail))
        if s_web:
            segments.append(("link", s_web, s_web.replace("https://", "").replace("http://", "")))
        if s_linkedin:
            segments.append(("link", s_linkedin, "LinkedIn"))
        if s_xing:
            segments.append(("link", s_xing, "Xing"))

        for i, seg in enumerate(segments):
            if i > 0:
                sep_run = fp.add_run("   |   ")
                sep_run.font.size = Pt(8.5)
                sep_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            if seg[0] == "text":
                r = fp.add_run(seg[1])
                r.font.size = Pt(8.5)
            else:
                _docx_add_hyperlink(fp, seg[1], seg[2], size_pt=8.5)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def create_cover_letter_pdf(letter: dict, sender: dict = None, recipient: dict = None,
                             date_str: str = "") -> BytesIO:
    """Erstellt dasselbe Anschreiben als sauber umbrochenes PDF (reportlab
    Platypus statt Zeile-fuer-Zeile-Canvas -> echter Fliesstext-Umbruch,
    saubere Seitenumbrueche bei laengeren Texten), inkl. grosser Betreffzeile
    und klickbarer Kontakt-Fusszeile."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_RIGHT, TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Frame, PageTemplate
    from xml.sax.saxutils import escape as _xml_escape

    sender = sender or {}
    recipient = recipient or {}

    s_phone = _cl_field(sender, "phone")
    s_mail = _cl_field(sender, "email")
    s_web = _cl_norm_url(_cl_field(sender, "website"))
    s_linkedin = _cl_norm_url(_cl_field(sender, "linkedin"))
    s_xing = _cl_norm_url(_cl_field(sender, "xing"))
    has_footer = bool(s_phone or s_mail or s_web or s_linkedin or s_xing)

    buf = BytesIO()
    bottom_margin = 2.6 * cm if has_footer else 2.0 * cm
    docpdf = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=2.5 * cm, bottomMargin=bottom_margin,
        leftMargin=2.5 * cm, rightMargin=2.0 * cm,
    )

    base_styles = getSampleStyleSheet()
    style_small = ParagraphStyle("clSmall", parent=base_styles["Normal"],
                                  fontName="Helvetica", fontSize=9.5, leading=13)
    style_normal = ParagraphStyle("clNormal", parent=base_styles["Normal"],
                                   fontName="Helvetica", fontSize=11, leading=17,
                                   spaceAfter=12)
    style_right = ParagraphStyle("clRight", parent=style_normal, alignment=TA_RIGHT,
                                  spaceAfter=16)
    # Betreffzeile: deutlich groesser + fett (vom Nutzer gewuenscht)
    style_subject = ParagraphStyle("clSubject", parent=style_normal, fontName="Helvetica-Bold",
                                    fontSize=14.5, leading=18, spaceBefore=6, spaceAfter=18)
    style_footer = ParagraphStyle("clFooter", parent=base_styles["Normal"], fontName="Helvetica",
                                   fontSize=8.5, leading=11, alignment=TA_CENTER,
                                   textColor=HexColor("#555555"))

    story = []

    def esc(t: str) -> str:
        return _xml_escape(t or "").replace("\n", "<br/>")

    s_name = _cl_field(sender, "name")
    s_addr = _cl_field(sender, "address")
    if s_name or s_addr or s_mail or s_phone:
        block = "<br/>".join(esc(x) for x in [s_name, s_addr, s_mail, s_phone] if x)
        story.append(Paragraph(block, style_small))
        story.append(Spacer(1, 18))

    story.append(Paragraph(esc(date_str), style_right))

    r_company = _cl_field(recipient, "company")
    r_att = _cl_field(recipient, "attention")
    r_addr = _cl_field(recipient, "address")
    if r_company or r_att or r_addr:
        block = "<br/>".join(esc(x) for x in [r_company, r_att, r_addr] if x)
        story.append(Paragraph(block, style_normal))
        story.append(Spacer(1, 12))

    subject = _cl_field(letter, "subject", "Bewerbung")
    story.append(Paragraph(esc(subject), style_subject))

    salutation = _cl_field(letter, "salutation", "Sehr geehrte Damen und Herren,")
    story.append(Paragraph(esc(salutation), style_normal))

    paragraphs = letter.get("paragraphs") or []
    if isinstance(paragraphs, str):
        paragraphs = [p.strip() for p in paragraphs.split("\n\n") if p.strip()]
    for para_text in paragraphs:
        story.append(Paragraph(esc(para_text), style_normal))

    closing = _cl_field(letter, "closing", "Mit freundlichen Gruessen")
    story.append(Spacer(1, 12))
    story.append(Paragraph(esc(closing), style_normal))

    sig_name = _cl_field(letter, "signature_name", s_name)
    if sig_name:
        story.append(Spacer(1, 36))
        story.append(Paragraph(esc(sig_name), style_normal))

    if has_footer:
        segs = []
        if s_phone:
            segs.append(esc(s_phone))
        if s_mail:
            segs.append(f'<link href="mailto:{esc(s_mail)}"><u>{esc(s_mail)}</u></link>')
        if s_web:
            web_label = esc(s_web.replace("https://", "").replace("http://", ""))
            segs.append(f'<link href="{esc(s_web)}"><u>{web_label}</u></link>')
        if s_linkedin:
            segs.append(f'<link href="{esc(s_linkedin)}"><u>LinkedIn</u></link>')
        if s_xing:
            segs.append(f'<link href="{esc(s_xing)}"><u>Xing</u></link>')
        footer_text = "   |   ".join(segs)

        def _draw_footer(canvas, doc_):
            canvas.saveState()
            canvas.setStrokeColor(HexColor("#CCCCCC"))
            canvas.setLineWidth(0.5)
            line_y = bottom_margin - 0.5 * cm
            canvas.line(2.5 * cm, line_y, A4[0] - 2.0 * cm, line_y)
            p = Paragraph(footer_text, style_footer)
            w, h = p.wrap(A4[0] - 4.5 * cm, bottom_margin)
            p.drawOn(canvas, 2.5 * cm, line_y - h - 4)
            canvas.restoreState()

        frame = Frame(2.5 * cm, bottom_margin, A4[0] - 4.5 * cm, A4[1] - 2.5 * cm - bottom_margin,
                       id="normal")
        docpdf.addPageTemplates([PageTemplate(id="withFooter", frames=[frame], onPage=_draw_footer)])

    docpdf.build(story)
    buf.seek(0)
    return buf
